from reportlab.pdfgen import canvas
from docx import Document
import subprocess
from subprocess import  Popen
import os

def convert_to_pdf(input_docx, out_folder):
    LIBRE_OFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()

def resumeGeneratorTXT(position, company, file):
    f=open(company + "CoverLetter.txt", "w+")
    g=open(file + ".txt","r")
    contents = g.read()
    f.write(ssentence + " " + position + " position with " + company + contents)
    f.close()
    g.close()
    return 0;

def resumeGeneratorDOC(position, company, file, salutation):
    document = Document()
    
    g=open(file+".txt","r")
    
    contents = g.read()
    
    p = document.add_paragraph(salutation)
    p = document.add_paragraph("I am excited to apply for your" + " " + position + " position with " + company +". " +contents)
    
    g.close()
    document.save(company + "CoverLetter.docx")
    
    return company + "CoverLetter.docx";

def resumeGeneratorPDF(position, company, file, salutation): #in progress
    docName = resumeGeneratorDOC(position, company, file, salutation)
    convert_to_pdf(docName,  os.getcwd())
    os.remove(docName)
    



